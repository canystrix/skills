"""Convert an interview-prep Markdown file to a Calibri-styled Word document.

Usage:
    python3 md_to_docx.py path/to/interview-prep-company.md [--out path.docx]

The converter handles the subset of Markdown used by the interview-prep
skill: #/##/### headings, blank-line paragraph breaks, ``**bold**`` and
``*italic*`` inline spans, ``-`` bullet lists, and ``---`` horizontal
rules (which are dropped because headings already provide structure).
"""

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor


FONT_BODY = "Calibri"
FONT_SIZE_BODY = 11
FONT_SIZE_H1 = 16
FONT_SIZE_H2 = 13
FONT_SIZE_H3 = 11
COLOR_DARK = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_SUBTLE = RGBColor(0x55, 0x55, 0x55)

INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.5):
    """Set page margins in centimetres."""
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def style_run(run, bold=False, italic=False, size=FONT_SIZE_BODY, color=COLOR_DARK):
    """Apply shared font styling to a run."""
    run.bold = bold
    run.italic = italic
    run.font.name = FONT_BODY
    run.font.size = Pt(size)
    run.font.color.rgb = color


def add_inline_runs(para, text, size=FONT_SIZE_BODY, color=COLOR_DARK):
    """Split text by bold/italic markers and add runs accordingly."""
    if not text:
        return
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            style_run(run, bold=True, size=size, color=color)
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            style_run(run, italic=True, size=size, color=color)
        else:
            run = para.add_run(part)
            style_run(run, size=size, color=color)


def add_heading(doc, text, level):
    """Add a heading paragraph with size/spacing tuned to the level."""
    sizes = {1: FONT_SIZE_H1, 2: FONT_SIZE_H2, 3: FONT_SIZE_H3}
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14 if level == 1 else 12 if level == 2 else 8)
    para.paragraph_format.space_after = Pt(6 if level <= 2 else 4)
    run = para.add_run(text)
    style_run(run, bold=True, size=sizes[level])


def add_bullet(doc, text):
    """Add a bulleted paragraph using Word's built-in list style."""
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    add_inline_runs(para, text)


def add_paragraph(doc, text, after=4):
    """Add a body paragraph with optional trailing space."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(after)
    add_inline_runs(para, text)


def convert(md_path, docx_path):
    """Render the Markdown file at md_path to a Word document at docx_path."""
    doc = Document()
    set_margins(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
        elif line.strip() == "---":
            continue
        elif line.lstrip().startswith("- "):
            add_bullet(doc, line.lstrip()[2:].strip())
        else:
            add_paragraph(doc, line.strip())

    doc.save(docx_path)


def main():
    """CLI entry point."""
    parser = argparse.ArgumentParser(
        description="Convert interview-prep markdown to a styled Word document."
    )
    parser.add_argument("md_path", type=Path, help="Input Markdown file")
    parser.add_argument(
        "--out",
        type=Path,
        default=None,
        help="Output .docx path (defaults to the input stem with .docx extension)",
    )
    args = parser.parse_args()

    out = args.out or args.md_path.with_suffix(".docx")
    convert(args.md_path, out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
